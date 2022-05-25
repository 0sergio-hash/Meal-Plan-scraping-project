
def get_syntax(): 
	print('''

	get_meals(when = 'next', email = True, 
	send_to = 'default@email.com')

	when : 'now' (today) or 'next' (next week)

	email : True or False

	send_to : The email address you wish to send the finished document to.

	''')


def get_meals(when = 'next', email = True, send_to = 'default@email.com'):
    
    #Library for inputting password
    import getpass

    #Library for fetching a webpage
    from urllib.request import urlopen as request

    #Library for parsing html from webpage
    from bs4 import BeautifulSoup as soup

    # Library to ask for username and passwords as input
    import pyinputplus as pyip 
    
    #Library for writing to a word document
    import docx
    from docx.shared import Pt #for adjusting font and size
    from docx.enum.text import WD_ALIGN_PARAGRAPH #for changing paragraph alignment
    from docx.shared import Inches, Cm #for changing page margains
    
    #Library to pull HTML content from Java enabled sites
    from selenium import webdriver
    
    #Needed to program in waiting for elements to appear before moving on with script
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    
    import re # For regular expressions. Will use to get date of meal plan
    
    import os #for saving docs 
    
    import smtplib #for sending emails
    from email.mime.multipart import MIMEMultipart
    from email.mime.text import MIMEText
    from email.mime.base import MIMEBase
    from email import encoders
    
    import time #add wait time to code
    
    # Status update
    print('STATUS: Dependencies imported.')
    
    # Depending on when parameter, get meal plan url for current of next week
    if when == 'next':
        url = 'https://www.strongrfastr.com/app/meal_plan/next' #save the url we want to open to variable 'url'
    elif when == 'now':
        url = 'https://www.strongrfastr.com/app/meal_plan/' #save the url we want to open to variable 'url'
    
    #Open a browser window
    driver = webdriver.Firefox()
    
    #Navigate to meal planning site
    driver.get(url)
    
    # Status update
    print('STATUS: Webpage open.')
    print()
    
    #Wait for the email login button to load
    try:
        element = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.ID, 'email-login-btn'))
        )
    except:
        driver.quit()
    
    #Define my email and password
    User_Name = pyip.inputStr(prompt = 'Enter email address: ')
    PWD = getpass.getpass()
    
   #Define login button, click button, pass in login credentials
    Login_button = driver.find_element_by_id('email-login-btn')
    Login_button.click()

    username = driver.find_element_by_id("email-field")
    password = driver.find_element_by_id("password-field")

    username.send_keys(User_Name)
    password.send_keys(PWD)
    
    #Wait for sign in button to load
    try:
        element = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.ID, 'submit-email-btn'))
        )
    except:
        driver.quit() 
    
    #Define sign in button, click to proceed to my profile and next week's meal plan
    submit_button = driver.find_element_by_id('submit-email-btn')
    submit_button.click()
    
    print()
    print('STATUS: Signed in.')
    
    #Wait for the root div to load (the element that contains the rest of the webpage)
    try:
        element = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.ID, 'root'))
        )
    except:
        driver.quit()
    
    #Use time to give page a couple seconds to load
    time.sleep(5) # Sleep for 5 seconds
    
    #Save page to variable
    source = driver.page_source
    
    #Use BeautifulSoup to parse the webpage
    page_soup = soup(source, "html.parser")
    
    #Save all meal cards to variable, save the first 4 (One day's meals) to a variable

    #object of all divs that contain a meal
    mydivs = page_soup.find_all("div", {"class": "meal-cards-day-wrap"}) #BeautifulSoup's findall() method returns a list
    
    meals = []

    if len(mydivs) > 0:
        meals = mydivs[0:4] 
        #assign the first four meals (they're the same all week) to variable 'meals'. This is the list we will iterate over
    else:
        print('ERROR:Page not loaded before scraped')
     
    #for each meal, extract the end of it's URL, append to a list
    
    links = []
    for meal in meals:
        for link in meal.find_all('a'):
            links.append(link)
    
    # Get the date for the meal plan in question
    
    link = links[0].get('href')

    dateRegex = re.compile(r'\d\d\d\d-\d\d-\d\d') #defines pattern to search for, creates a regular expression object
    
    date = dateRegex.search(link) #searchses for pattern and returns a match object to a variable
    
    date_text = date.group() # the group() method returns a string we can save to the date_text variable
    
    # Status Update
    print('STATUS: Meals found.')
    
    #Change current working directory

    os.chdir('C:\\Users\\examplepath\\meal_plans')


    #Create a word document 

    doc = docx.Document()

    # Set to narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    date_p = doc.add_paragraph(date_text)

    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Write meal titles and dishes

    # Status list to iterate over
    status = ['25%','50%','75%','100%']
    for meal in meals: #for each meal in the meals list
        title_p = doc.add_paragraph()

        title_p.add_run(meal.div.text + ': ')

        main_dishes = meal.find_all('div',{'class':'meal-plan-recipe-title'}) #create dictionary of all main dishes
        side_dishes = meal.find_all('div',{'class':'low-emphasis lh1 abbr-text'}) #create dictionary of all side dishes
        for dish in main_dishes:

            title_p.add_run(dish.text)

        for side in side_dishes:

            title_p.add_run(side.text)

        for run in title_p.runs:
            font = run.font

            font.name = 'Calibri'
            font.size = Pt(12)
            font.underline = True
            font.all_caps = True


        # Write all ingredients and quantities

        for link in meal.find_all('a'):
            meal_url = 'https://www.strongrfastr.com' + link.get('href')
            driver.get(meal_url) #navigate to the URL for the meal on each iteration of the loop

            #Wait for the root div to load (the element that contains the rest of the webpage)
            try:
                element = WebDriverWait(driver, 20).until(
                    EC.presence_of_element_located((By.ID, 'root'))
                )
            except:
                driver.quit()

            #Use time to give page a couple seconds to load
            time.sleep(5) # Sleep for 5 seconds

            #Save page to variable
            meal_source = driver.page_source

            #Use BeautifulSoup to parse the webpage
            meal_page_soup = soup(meal_source, "html.parser")

            #BeautifulSoup's findall() method returns a list
            #object of all divs that contain a dish, and another of side dishes

            dish_divs = meal_page_soup.find_all("div", {"class": "mec-wrapper mb10"}) 
            side_dish_divs = meal_page_soup.find_all("div", {"class": "side-dish-container"})  


            # Write all main dishes
            for dish in dish_divs: # for every dish in the list of main dishes
                for link in dish.find_all('a'): #find all the links
                    main_url = 'https://www.strongrfastr.com' + link.get('href') #assign url to variable
                    driver.get(main_url) #navigate to the URL for the dish on each iteration of the loop

                    #Wait for the root div to load (the element that contains the rest of the webpage)
                    try:
                        element = WebDriverWait(driver, 20).until(
                            EC.presence_of_element_located((By.ID, 'root'))
                        )
                    except:
                        driver.quit()

                    #Use time to give page a couple seconds to load
                    time.sleep(5) # Sleep for 5 seconds

                    #Save page to variable
                    main_dish_source = driver.page_source

                    #Use BeautifulSoup to parse the webpage
                    main_dish_soup = soup(main_dish_source, "html.parser")


                    #Scrape dish names
                    dish_name_div = main_dish_soup.find("div", {"class": "max-list ml-center mt0"}) 
                    dish_name = dish_name_div.find('h3').text

                    #Scrape ingredients

                    main_ingredients_list = main_dish_soup.find_all("div", {"class": "basic-heading"})

                    #scrape quantities

                    main_quantities_list = main_dish_soup.find_all("div", {"class": "basic-heading-sub"})[1:]
                    # the 0th index returns nutrition info because its tag is a partial match. So we exclude it

                    main_p = doc.add_paragraph()

                    main_title = main_p.add_run(dish_name + ': ')

                    font = main_title.font

                    font.name = 'Calibri'
                    font.size = Pt(12)
                    font.underline = True
                    font.all_caps = True

                    for i in range(len(main_ingredients_list)):

                        main_p.add_run(main_ingredients_list[i].text + ' ')

                        main_p.add_run(main_quantities_list[i].text + ' ')

                    for run in main_p.runs[1:]:
                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(12)


            #Repeat for side dishes
            for side_dish in side_dish_divs: # for every dish in the list of main dishes
                for side_link in side_dish.find_all('a'): #find all the links
                    side_url = 'https://www.strongrfastr.com' + side_link.get('href') #assign url to variable
                    driver.get(side_url) #navigate to the URL for the dish on each iteration of the loop

                    #Wait for the root div to load (the element that contains the rest of the webpage)
                    try:
                        element = WebDriverWait(driver, 20).until(
                            EC.presence_of_element_located((By.ID, 'root'))
                        )
                    except:
                        driver.quit()

                    #Use time to give page a couple seconds to load
                    time.sleep(5) # Sleep for 5 seconds

                    #Save page to variable
                    side_dish_source = driver.page_source

                    #Use BeautifulSoup to parse the webpage
                    side_dish_soup = soup(side_dish_source, "html.parser")


                    #Scrape dish names
                    side_name_div = side_dish_soup.find("div", {"class": "max-list ml-center mt0"}) 
                    side_name = side_name_div.find('h3').text

                    #Scrape ingredients

                    side_ingredients_list = side_dish_soup.find_all("div", {"class": "basic-heading"})

                    #scrape quantities

                    side_quantities_list = side_dish_soup.find_all("div", {"class": "basic-heading-sub"})[1:]
                    # the 0th index returns nutrition info because its tag is a partial match. So we exclude it

                    side_p = doc.add_paragraph()

                    side_title = side_p.add_run(side_name + ': ')

                    font = side_title.font

                    font.name = 'Calibri'
                    font.size = Pt(12)
                    font.underline = True
                    font.all_caps = True

                    side_ingredients_p = doc.add_paragraph()
                    for s in range(len(side_ingredients_list)):

                        side_p.add_run(side_ingredients_list[s].text + ' ')

                        side_p.add_run(side_quantities_list[s].text + ' ')


                    for run in side_p.runs[1:]:

                        font = run.font
                        font.name = 'Calibri'
                        font.size = Pt(12)
        
        # Status update
        index = meals.index(meal)
        print('STATUS: Meals scraped: ' + status[index])

        # Update index
        #index += 1   
            
    #Close the browser window
    driver.close()

    #Save our word document
    try:
        doc.save('Comidas ' + date_text +'.docx')
    except:
        print('File save error')
    
    # Status update
    print('STATUS: Meal plan document saved.')
    print()

    # Finally, send our new document to the specified email address
    #Istructions on https://www.geeksforgeeks.org/send-mail-attachment-gmail-account-using-python/
    
    if email == True:
        
        fromaddr = 'default@email.com'
        toaddr = send_to

        # instance of MIMEMultipart
        msg = MIMEMultipart()

        # storing the senders email address  
        msg['From'] = fromaddr

        # storing the receivers email address 
        msg['To'] = toaddr

        # storing the subject 
        msg['Subject'] = "Comidas"

        # string to store the body of the mail
        body = " "

        # attach the body with the msg instance
        msg.attach(MIMEText(body, 'plain'))

        # open the file to be sent 
        filename = 'Comidas ' + date_text +'.docx'
        attachment = open('C:\\Users\\examplepath\\meal_plans\\'+'Comidas ' + date_text +'.docx', "rb")

        # instance of MIMEBase and named as p
        p = MIMEBase('application', 'octet-stream')

        # To change the payload into encoded form
        p.set_payload((attachment).read())

        # encode into base64
        encoders.encode_base64(p)

        p.add_header('Content-Disposition', "attachment; filename= %s" % filename)

        # attach the instance 'p' to instance 'msg'
        msg.attach(p)

        # creates SMTP session
        s = smtplib.SMTP('smtp.gmail.com', 587)

        # start TLS for security
        s.starttls()

        # Authentication
        email_password = getpass.getpass(prompt = 'Email Password: ')
        s.login(fromaddr, email_password)

        # Converts the Multipart msg into a string
        text = msg.as_string()

        # sending the mail
        s.sendmail(fromaddr, toaddr, text)

        # terminating the session
        s.quit()
        
        # Status update
        print()
        print('STATUS: Email sent.')
